import docx

doc = docx.Document('demo.docx')
# print(doc.paragraphs[1].text)
# for run in doc.paragraphs[1].runs:
    # print(run.text)        # read text of every run object

# fullText = []
# for para in doc.paragraphs:
#    fullText.append(para.text)
# print("\n".join(fullText))      # read full text of docx
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True

doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')

doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)

doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('restyled.docx')
